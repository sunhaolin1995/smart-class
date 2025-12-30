from docx import Document
import sys
import os

# Import the function from app.py (assuming it's in the same dir)
sys.path.append("/Users/haolinsun/Desktop/agentProject/lesson plan generator")
from app import get_table_structure

def create_test_doc():
    doc = Document()
    
    # CASE 1: Key | Value (Empty)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Course Name"
    # cell(0, 1) is empty -> Should match
    
    # CASE 2: Key (Row 1) -> Value (Row 2, Empty)
    table2 = doc.add_table(rows=2, cols=1)
    table2.cell(0, 0).text = "Teaching Objectives"
    # cell(1, 0) is empty -> Should match
    
    # CASE 3: False Positive Check
    # Text | Text -> No match
    table3 = doc.add_table(rows=1, cols=2)
    table3.cell(0, 0).text = "Existing Key"
    table3.cell(0, 1).text = "Existing Value"
    
    return doc

def test_parsing():
    print("Creating test document...")
    doc = create_test_doc()
    
    print("Running parser...")
    structure = get_table_structure(doc)
    
    print(f"Found {len(structure)} fields.")
    for item in structure:
        print(f"  - Key: '{item['key_text']}' -> Target: {item['target_coords']}")
        
    # Assertions
    keys = [item['key_text'] for item in structure]
    assert "Course Name" in keys, "Failed to find 'Course Name' (Right mapping)"
    assert "Teaching Objectives" in keys, "Failed to find 'Teaching Objectives' (Down mapping)"
    assert "Existing Key" not in keys, "Incorrectly mapped 'Existing Key' which already has a value"
    
    print("\n✅ Parsing logic test passed!")

if __name__ == "__main__":
    try:
        test_parsing()
    except AssertionError as e:
        print(f"\n❌ Test Failed: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        sys.exit(1)
