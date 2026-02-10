import streamlit as st
import os
from docx import Document
from docx.shared import Pt
import json
import time
import re
import math
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from io import BytesIO

# --- Configuration ---
st.set_page_config(page_title="AI 智能教案生成器 (V16 Flagship)", layout="wide", initial_sidebar_state="expanded")

# --- UI Components: Console Logger ---
class ConsoleLogger:
    def __init__(self):
        self.container = st.empty()
        self.logs = []

    def log(self, message, icon="🤖"):
        timestamp = time.strftime("%H:%M:%S")
        self.logs.append(f"`{timestamp}` {icon} {message}")
        with self.container.container():
            with st.expander("🖥️ AI 运行终端 (实时日志)", expanded=True):
                for log in self.logs[-5:]: # Show last 5 logs
                    st.markdown(log)
    
    def clear(self):
        self.container.empty()
        self.logs = []

# --- Logic: Helper Functions ---
def set_cell_text_preserving_style(cell, text):
    """保留原有格式写入文本"""
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return

    paragraph = cell.paragraphs[0]
    style_run = paragraph.runs[0] if paragraph.runs else None
    
    paragraph.clear()
    run = paragraph.add_run(text)
    
    if style_run:
        run.bold = style_run.bold
        run.italic = style_run.italic
        run.font.name = style_run.font.name
        if style_run.font.size:
            run.font.size = style_run.font.size

def extract_json_safe(content):
    """
    JSON 提取与修复 (V14+)
    """
    if "```json" in content:
        content = content.split("```json")[1].split("```")[0]
    elif "```" in content:
        content = content.split("```")[1].split("```")[0]
    
    content = content.strip()
    # 修复尾部逗号
    content = re.sub(r',(\s*})', r'\1', content)
    content = re.sub(r',(\s*])', r'\1', content)
    
    try:
        return json.loads(content)
    except Exception:
        return None

# --- Logic: Structure Parsing (V16 Optimized) ---
def get_table_structure(doc, logger=None):
    """
    【V16 结构解析引擎】
    1. 智能处理跨页重复表头（忽略重复的红色“课中”）。
    2. 确保抓取到“课后”和“巩固拓展”。
    3. 全局唯一行号 Key。
    """
    if logger: logger.log("正在扫描文档结构 (V16 智能版)...", "🔍")
    
    structure = []
    processed_cell_ids = set() 
    processed_keys = set()     

    # 只要不是死循环，尽可能多抓，由 Prompt 控制内容
    MAX_ROWS_PER_PHASE = 100 

    def is_instructional(text):
        return len(text) > 30 or any(k in text for k in ["思政案例", "确保思政", "比例可根据"])

    for t_idx, table in enumerate(doc.tables):
        rows = table.rows
        all_text = "".join([c.text for r in rows for c in r.cells])
        
        # --- 策略 A：教学过程矩阵表 ---
        if any(k in all_text for k in ["教师活动", "学生活动", "设计意图"]):
            if logger: logger.log(f"正在解析教学过程表...", "🎯")
            
            col_map = {}
            # 扫描前几行找列名
            for r_idx in range(min(5, len(rows))):
                for c_idx in range(len(table.columns)):
                    txt = table.cell(r_idx, c_idx).text.strip()
                    if txt in ["教学环节", "教学内容", "教师活动", "学生活动", "设计意图"]:
                        col_map[c_idx] = txt

            current_phase = "教学过程"
            # 使用列表来记录已处理的阶段，用于判断是否重复
            seen_phases = []

            for r in range(len(rows)):
                # 获取该行纯文本，用于判断阶段
                row_raw_text = "".join(list(dict.fromkeys([c.text.strip() for c in rows[r].cells])))
                
                # 1. 识别阶段切换
                if row_raw_text in ["课前", "课中", "课后", "巩固拓展"]:
                    # 【核心修改】：如果这个阶段名和当前阶段一样，说明是跨页重复表头，直接忽略
                    if row_raw_text == current_phase:
                        continue 
                    
                    # 这是一个新的阶段
                    current_phase = row_raw_text
                    seen_phases.append(current_phase)
                    continue
                
                # 2. 正常抓取填空点
                for c_idx, col_name in col_map.items():
                    target_cell = table.cell(r, c_idx)
                    
                    # 排除非空格子、排除表头本身
                    if not target_cell.text.strip() and target_cell.text.strip() != col_name:
                        if target_cell._tc not in processed_cell_ids:
                            # 构造唯一 Key：阶段 > 标题 > 行号
                            full_key = f"{current_phase} > {col_name}_行{r}"
                            
                            structure.append({
                                'key_text': full_key,
                                'original_text': col_name,
                                'target_coords': (t_idx, r, c_idx),
                                'is_teaching_process': True
                            })
                            processed_cell_ids.add(target_cell._tc)
            continue

        # --- 策略 B：通用信息表 ---
        for r in range(len(rows)):
            for c in range(len(table.columns)):
                cell = table.cell(r, c)
                text = cell.text.strip().replace("\n", "").replace(" ", "")
                
                if not text or is_instructional(text): continue
                if text in processed_keys and len(text) < 10: continue

                target = None
                if c + 1 < len(table.columns):
                    r_c = table.cell(r, c + 1)
                    if not r_c.text.strip(): target = (r, c + 1, r_c._tc)
                if not target and r + 1 < len(rows):
                    d_c = table.cell(r + 1, c)
                    if not d_c.text.strip(): target = (r + 1, c, d_c._tc)
                
                if target:
                    tr, tc, t_id = target
                    if t_id not in processed_cell_ids:
                        full_key = text
                        p_header = table.cell(r, 0).text.strip()
                        if p_header in ["学情分析", "教学目标", "教学资源", "教学反思"]:
                            if p_header != text: full_key = f"{p_header} > {text}"

                        structure.append({
                            'key_text': full_key,
                            'original_text': text,
                            'target_coords': (t_idx, tr, tc),
                            'is_teaching_process': False
                        })
                        processed_cell_ids.add(t_id)
                        processed_keys.add(text)

    return structure

# --- Logic: Chunked Generation Engine (V16 Optimized) ---
def generate_deep_content_chunked(user_inputs, doc_keys, api_key, logger):
    """
    【V16 差异化生成引擎】
    1. Batch Size 提升至 45，大幅减少请求次数。
    2. Prompt 区分对待：学情/目标要详实，过程要干练且无编号。
    """
    llm = ChatOpenAI(
        model="deepseek-chat", 
        temperature=0.7, # 稍微回升温度，让长文写得更好
        base_url="https://api.deepseek.com",
        openai_api_key=api_key
    )
    
    all_keys = [item['key_text'] for item in doc_keys]
    
    # 【修改点】：增加 Batch Size 到 45，减少分组数量
    BATCH_SIZE = 45
    
    total_batches = math.ceil(len(all_keys) / BATCH_SIZE)
    final_mapping = {}
    
    logger.log(f"任务总量: {len(all_keys)} 个字段，合并为 {total_batches} 批次极速生成...", "🚀")
    
    progress_bar = st.progress(0)
    
    for i in range(total_batches):
        start_idx = i * BATCH_SIZE
        end_idx = start_idx + BATCH_SIZE
        current_batch_keys = all_keys[start_idx:end_idx]
        
        logger.log(f"正在生成第 {i+1}/{total_batches} 批...", "⏳")
        
        # --- 核心修改：Prompt 差异化约束 ---
        system_prompt = """
你是一位顶尖的教案设计专家。请根据课程背景，填写教案空格。

## ⚠️ 核心指令：差异化写作风格 (Differentiated Style)

请根据 **Key 的类型** 自动切换写作模式：

### 模式 A：【学情分析】与【教学目标】类
-   **适用 Key**：包含 "学情"、"目标"、"基础"、"分析" 的字段。
-   **要求**：**内容详实、具体**。可以写 100 字左右，分点阐述，深入分析学生特点和教学目的。

### 模式 B：【教学过程】类 (表格内容)
-   **适用 Key**：包含 "课前"、"课中"、"课后"、"活动"、"内容" 的字段。
-   **要求**：**短小精悍**。
-   **❌ 严禁使用编号**：禁止使用 "1. 2. 3." 或 "- " 列表符号。**直接写动作！**
-   **示例**：
    -   ❌ 错误：1. 教师播放视频。2. 提问学生。
    -   ✅ 正确：播放行业应用视频，提问引发思考，展示代码运行效果。

### 模式 C：【思政与解决措施】
-   **要求**：必须结合具体知识点，拒绝空话。

## 格式铁律
-   输出合法的 JSON：`{{ "Key": "Value" }}`
-   Key 必须用双引号。
-   **严禁**尾部逗号。
"""
        
        human_template = """
【课程背景】: {user_inputs_json}

【本次需填写的 Key】: 
{batch_keys_json}

请严格按照“差异化风格”填充上述 Key。
"""
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", human_template)
        ])
        
        chain = prompt | llm
        
        retry_count = 0
        success = False
        
        while retry_count < 2 and not success:
            try:
                response = chain.invoke({
                    "user_inputs_json": json.dumps(user_inputs, ensure_ascii=False),
                    "batch_keys_json": json.dumps(current_batch_keys, ensure_ascii=False)
                })
                
                batch_result = extract_json_safe(response.content)
                
                if batch_result:
                    final_mapping.update(batch_result)
                    success = True
                else:
                    logger.log(f"第 {i+1} 批次 JSON 解析失败，重试中...", "⚠️")
                    retry_count += 1
                    
            except Exception as e:
                logger.log(f"API 请求失败: {e}，冷却后重试...", "⚠️")
                retry_count += 1
                time.sleep(1) 
        
        if not success:
            logger.log(f"第 {i+1} 批次失败，已跳过。", "❌")
        
        progress_bar.progress((i + 1) / total_batches)

    # 硬逻辑补丁
    logger.log("生成完毕，正在整合数据...", "🧩")
    
    manual_overrides = {
        "授课时间": user_inputs.get("时间"),
        "教案序号 > 授课时间": user_inputs.get("时间"),
        "授课地点": user_inputs.get("地点"),
        "教案序号 > 授课地点": user_inputs.get("地点"),
        "授课班级": user_inputs.get("班级"),
        "授课内容 > 授课班级": user_inputs.get("班级"),
        "教师姓名": user_inputs.get("教师姓名")
    }
    
    for k, v in manual_overrides.items():
        if v:
            final_mapping[k] = v
            
    return final_mapping

# --- Main App ---

def main():
    st.markdown("## 🤖 AI 智能教案生成器 (V16 Flagship)")
    
    # 0. Global Logger
    logger = ConsoleLogger()

    # 1. Sidebar Config
    with st.sidebar:
        st.header("⚙️ 1. 基础配置")
        api_key = st.text_input("DeepSeek API Key", type="password")
        
        st.header("📝 2. 课程基础信息")
        
        col1, col2 = st.columns(2)
        serial_no = col1.text_input("教案序号", "No. 01")
        time_val = col2.text_input("授课时间", "2024-03-20")

        dept = st.text_input("部门/院系", "信息工程学院")
        teacher = st.text_input("教师姓名", "张三")
        
        course_type = st.selectbox("课程性质 (AI可覆盖)", ["理论课", "实践课", "理实一体化", "研讨课"])
        
        user_inputs = {
            "教案序号": serial_no,
            "时间": time_val,
            "部门": dept,
            "教师姓名": teacher,
            "课程性质": course_type
        }

        with st.expander("📚 更多课程细节 (选填)", expanded=False):
            user_inputs["课程名称"] = st.text_input("课程名称", "Python 程序设计")
            user_inputs["班级"] = st.text_input("班级", "23级计算机1班")
            user_inputs["地点"] = st.text_input("授课地点", "A305")
            user_inputs["授课学时"] = st.number_input("学时", 1, 4, 2)
            user_inputs["授课形式"] = st.selectbox("授课形式", ["线下面授", "线上直播", "混合式教学"])
            user_inputs["使用教材"] = st.text_input("使用教材", "《Python编程：从入门到实践》")
            user_inputs["考核方式"] = st.selectbox("考核方式", ["考查", "考试", "过程化考核"])

        st.header("🧠 3. 核心内容输入")
        topic_outline = st.text_area("本节课主题 & 大纲", height=250, 
                                     placeholder="输入本节课的主题，例如：\n主题：Python 循环结构\n1. while 循环\n2. fo 循环\n3. 案例实战")
        user_inputs["课程大纲"] = topic_outline

    # 2. Main Area
    uploaded_file = st.file_uploader("📂 上传 Word 教案模板 (.docx)", type=["docx"])

    if uploaded_file and st.button("🚀 开始生成", type="primary"):
        if not api_key:
            st.error("请先在左侧输入 DeepSeek API Key")
            return
        
        if not topic_outline:
            st.warning("请填写【课程主题 & 大纲】，否则 AI 无法生成内容。")
            return

        # Step 1: Parse
        doc = Document(uploaded_file)
        structure = get_table_structure(doc, logger)
        
        if not structure:
            st.warning("未能识别到表格结构。请确保文档包含标准表格。")
            return

        # Step 2: Generate (V16)
        mapping = generate_deep_content_chunked(user_inputs, structure, api_key, logger)
        
        # Step 3: Fill
        if mapping:
            logger.log("正在将内容写入文档...", "💾")
            fill_count = 0
            
            # Progress bar for filling
            my_bar = st.progress(0)
            total_items = len(structure)
            
            for i, item in enumerate(structure):
                key = item['key_text']
                target_coords = item['target_coords']
                original_text = item['original_text']
                
                content = mapping.get(key) or mapping.get(original_text)
                
                if content:
                    t_idx, r, c = target_coords
                    target_cell = doc.tables[t_idx].cell(r, c)
                    set_cell_text_preserving_style(target_cell, str(content))
                    fill_count += 1
                    if i % 10 == 0: 
                        logger.log(f"已填入: {key} -> {str(content)[:10]}...", "📝")
                
                my_bar.progress(min((i + 1) / total_items, 1.0))

            logger.log(f"🎉 全部完成！共填充 {fill_count} 个字段。", "✅")
            st.success(f"生成成功！")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="⬇️ 下载生成的教案",
                data=buffer,
                file_name="generated_lesson_plan_v16.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
if __name__ == "__main__":
    main()